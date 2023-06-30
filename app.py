from flask import Flask
from flask import render_template, request, redirect, session
from flaskext.mysql import MySQL
from datetime import datetime
from flask import send_from_directory
import os
from flask import current_app
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
import win32api

app=Flask(__name__)
app.secret_key="celumax"
mysql=MySQL()


app.config["MYSQL_DATABASE_HOST"]="localhost"   #dominio, localhost, entre otros 
app.config["MYSQL_DATABASE_USER"]="root"
app.config["MYSQL_DATABASE_PASSWORD"]=""
app.config["MYSQL_DATABASE_DB"]="sitio"
mysql.init_app(app)


@app.route("/")
def inicio():
    return render_template("sitio/index.html")

@app.route("/img/<imagen>")
def imagenes(imagen):
    print(imagen)
    return send_from_directory(os.path.join("templates/sitio/img"), imagen)

@app.route("/admin/recibos")
def recibos():
    if not "login" in session:
        return redirect("/admin/login")
    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute("SELECT * FROM `recibos`")
    recibos=cursor.fetchall()
    conexion.commit()
    print(recibos)

    return render_template("admin/lrecibos.html", recibos=recibos)

@app.route("/productos")
def productos():

    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute("SELECT * FROM `productos`")
    productos=cursor.fetchall()
    conexion.commit()
    print(productos)

    return render_template("sitio/nproductos.html", productos=productos)

@app.route("/admin/adm")
def adm():

    if not "login" in session:
        return redirect("/admin/login")
    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute("SELECT * FROM `productos`")
    productos=cursor.fetchall()
    conexion.commit()
    print(productos)

    return render_template("admin/oadmin.html", productos=productos)

@app.route("/admin")
def admin_index():
    if not "login" in session:
        return redirect("/admin/login")
    return render_template("admin/index.html")

@app.route("/admin/login")
def admin_login():
    return render_template("admin/login.html")

@app.route("/admin/login", methods=["POST"])
def admin_login_post():
    _usuario=request.form["txtUsuario"]
    _password=request.form["txtPassword"]
    print(_usuario)
    print(_password)

    if _usuario=="Celumax2023" and _password=="Condor2020":
        session["login"]=True
        session["usuario"]="Celumax"
        return redirect("/admin") 

    return render_template("admin/login.html")

@app.route("/admin/cerrar")
def admin_login_cerrar():
    session.clear()
    return redirect("/admin/login")

@app.route("/admin/adm/guardar", methods=["POST"])
def admin_producto_guardar():
    if not "login" in session:
        return redirect("/admin/login")

    _producto=request.form["txtProducto"]
    _img=request.files["txtImagen"]
    _precio=request.form["txtPrecio"]
    _no=request.form["txtNumber"]

    tiempo  = datetime.now()
    horaActual= tiempo.strftime("%Y%H%M%S")

    if _img.filename !="":
        nuevoNombre=horaActual+"_"+_img.filename
        _img.save("templates/sitio/img/" +nuevoNombre)

    sql="INSERT INTO `productos` (`ID`, `Producto`, `Imagen`, `Precio`, `No`) VALUES (NULL,%s, %s, %s,%s);"
    datos2=(_producto,nuevoNombre,_precio,_no)
    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute(sql,datos2)
    conexion.commit()

    print(_producto)
    print(_img)
    print(_precio)
    print(_no)

    return redirect("/admin/adm")

@app.route("/admin/adm/borrar", methods=["POST"])
def admin_producto_borrar():
    if not "login" in session:
        return redirect("/admin/login")


    _id2=request.form["txtID2"]
    print(_id2)

    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute("SELECT imagen FROM `productos` WHERE id=%s",(_id2))
    producto=cursor.fetchall()
    conexion.commit()
    print(producto)

    if os.path.exists("templates/sitio/img/"+str(producto[0][0])):
        os.unlink("templates/sitio/img/"+str(producto[0][0]))

    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute("DELETE FROM productos WHERE id=%s",(_id2))
    conexion.commit()

    return redirect("/admin/adm")

@app.route("/admin/recibos/guardar", methods=["POST"])
def admin_recibos_guardar():
    if not "login" in session:
        return redirect("/admin/login")
    _nombre=request.form["txtNombrerecibo"]
    _cc=request.form["txtCC"]
    _tel=request.form["txtTel"]
    _equipo=request.form["txtEquipo"]
    _imei=request.form["txtImei"]
    _procedimiento=request.form["txtProcedimiento"]
    _valor=request.form["txtValor"]
    _abono=request.form["txtAbono"]
    _clave=request.form["txtClave"]


    sql="INSERT INTO `recibos` (`ID`, `Fecha`, `Nombre`, `CC.`, `Tel.`, `Equipo`, `Imei`, `Procedimiento`, `Valor`, `Abono`, `Clave`) VALUES (NULL, current_timestamp(), %s, %s, %s,%s,%s, %s, %s, %s, %s);"
    datos=(_nombre,_cc,_tel,_equipo,_imei,_procedimiento,_valor,_abono,_clave)
    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute(sql,datos)
    conexion.commit()

    print(_nombre)
    print(_cc)
    print(_tel)
    print(_equipo)
    print(_imei)
    print(_procedimiento)
    print(_valor)
    print(_abono)
    print(_clave)
    return redirect("/admin/recibos")

@app.route("/admin/recibos/borrar", methods=["POST"])
def admin_recibo_borrar():
    if not "login" in session:
        return redirect("/admin/login")
    _id=request.form["txtID"]
    print(_id)

    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute("SELECT * FROM `recibos` WHERE id=%s",(_id))
    recibos=cursor.fetchall()
    conexion.commit()
    print(recibos)

    conexion=mysql.connect()
    cursor=conexion.cursor()
    cursor.execute("DELETE FROM recibos WHERE id=%s",(_id))
    conexion.commit()

    return redirect("/admin/recibos")

@app.route("/admin/recibos/imprimir", methods=["POST"])
def admin_recibo_imprimir():
    if not "login" in session:
        return redirect("/admin/login")

    _nombre = request.form["txtNombrerecibo"]
    _cc = request.form["txtCC"]
    _tel = request.form["txtTel"]
    _equipo = request.form["txtEquipo"]
    _imei = request.form["txtImei"]
    _procedimiento = request.form["txtProcedimiento"]
    _valor = request.form["txtValor"]
    _abono = request.form["txtAbono"]
    _clave = request.form["txtClave"]

    datos = {
        "Nombre": _nombre,
        "CC": _cc,
        "Tel": _tel,
        "Equipo": _equipo,
        "Imei": _imei,
        "Procedimiento": _procedimiento,
        "Valor": _valor,
        "Abono": _abono,
        "Clave": _clave
    }

    # Construir la ruta al documento existente
    ruta_documento = os.path.join(current_app.root_path, "templates", "sitio", "recibos", "Recibos.docx")

    # Abrir el documento existente
    document = Document(ruta_documento)

    # Obtener el cuerpo principal del documento
    cuerpo_principal = document.paragraphs

    # Crear un nuevo párrafo para agregar los inputs
    nuevo_parrafo = cuerpo_principal[11].insert_paragraph_before()
    nuevo_parrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    nuevo_parrafo.space_after = Pt(12)

    # Añadir los inputs al párrafo
    for campo, valor in datos.items():
        run_campo = nuevo_parrafo.add_run(f"{campo}:")
        run_campo.bold = True
        run_campo.font.size = Pt(9)
        
        run_valor = nuevo_parrafo.add_run(f" {valor}")
        run_valor.font.size = Pt(9)

        run_valor.add_break()

    # Guardar el documento modificado
    ruta_documento_modificado = os.path.join(current_app.root_path, "templates", "sitio", "recibos", "Recibos", f"{_nombre}.docx")
    document.save(ruta_documento_modificado)

    # Imprimir el documento
    win32api.ShellExecute(0, "print", ruta_documento_modificado, None, ".", 0)

    return redirect("/admin/recibos")

if __name__ == "__main__":
    app.run(debug=True)